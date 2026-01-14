"""
Resume Parser - Extract text and sections from various file formats
"""

from pathlib import Path
from typing import Dict, Optional, List
import logging
import re

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Parse resumes from various file formats.
    Supports: PDF, DOCX, DOC, TXT
    """

    # Common section headers for detection
    SECTION_PATTERNS = {
        "contact": r"(?i)(contact\s*info|contact\s*details|personal\s*info)",
        "summary": r"(?i)(summary|profile|objective|about\s*me|professional\s*summary)",
        "experience": r"(?i)(experience|work\s*history|employment|professional\s*experience|career\s*history)",
        "education": r"(?i)(education|academic|qualifications|degrees)",
        "skills": r"(?i)(skills|technical\s*skills|core\s*competencies|expertise|technologies)",
        "projects": r"(?i)(projects|portfolio|key\s*projects)",
        "certifications": r"(?i)(certifications|certificates|licenses|credentials)",
        "awards": r"(?i)(awards|honors|achievements|recognition)",
        "publications": r"(?i)(publications|papers|research)",
        "languages": r"(?i)(languages|language\s*skills)",
        "interests": r"(?i)(interests|hobbies|activities)"
    }

    async def parse(self, file_path: Path) -> Dict:
        """
        Parse a resume file and extract content.

        Args:
            file_path: Path to the resume file

        Returns:
            Dict containing:
                - text: Full extracted text
                - sections: Detected sections with content
                - metadata: File metadata
        """
        file_ext = file_path.suffix.lower()

        # Extract text based on file type
        if file_ext == ".pdf":
            text = await self._parse_pdf(file_path)
        elif file_ext == ".docx":
            text = await self._parse_docx(file_path)
        elif file_ext == ".doc":
            text = await self._parse_doc(file_path)
        elif file_ext == ".txt":
            text = await self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Clean and normalize text
        text = self._clean_text(text)

        # Detect sections
        sections = self._detect_sections(text)

        # Extract contact info
        contact_info = self._extract_contact_info(text)

        return {
            "text": text,
            "sections": sections,
            "contact_info": contact_info,
            "metadata": {
                "filename": file_path.name,
                "file_type": file_ext,
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        }

    async def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            return "\n".join(text_parts)

        except ImportError:
            # Fallback to PyPDF2
            try:
                from PyPDF2 import PdfReader

                text_parts = []
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text_parts.append(page.extract_text())

                return "\n".join(text_parts)

            except ImportError:
                logger.error("No PDF library available. Install pdfplumber or PyPDF2")
                raise ImportError("PDF parsing requires pdfplumber or PyPDF2")

    async def _parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            return "\n".join(text_parts)

        except ImportError:
            logger.error("python-docx not installed")
            raise ImportError("DOCX parsing requires python-docx")

    async def _parse_doc(self, file_path: Path) -> str:
        """Extract text from legacy DOC file (pre-2007 Word format)"""
        # Try antiword first (most reliable for .doc)
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Try catdoc as fallback
        try:
            import subprocess
            result = subprocess.run(
                ['catdoc', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Try using olefile for basic extraction
        try:
            import olefile

            if not olefile.isOleFile(file_path):
                raise ValueError("Not a valid DOC file")

            ole = olefile.OleFileIO(file_path)

            # Try to extract from WordDocument stream
            if ole.exists('WordDocument'):
                # For complex .doc files, we extract what text we can
                text_parts = []

                # Try to read any text streams
                for stream in ole.listdir():
                    stream_path = '/'.join(stream)
                    try:
                        if any(name in stream_path.lower() for name in ['text', 'content', 'document']):
                            data = ole.openstream(stream).read()
                            # Try to decode as text
                            try:
                                text = data.decode('utf-16-le', errors='ignore')
                                # Filter printable characters
                                text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
                                if text.strip():
                                    text_parts.append(text)
                            except:
                                pass
                    except:
                        continue

                ole.close()

                if text_parts:
                    return '\n'.join(text_parts)

            ole.close()
            raise ValueError("Could not extract text from DOC file")

        except ImportError:
            pass

        # Last resort: try to read as binary and extract printable text
        try:
            with open(file_path, 'rb') as f:
                content = f.read()

            # Try UTF-16 LE decoding (common in .doc files)
            try:
                text = content.decode('utf-16-le', errors='ignore')
            except:
                text = content.decode('latin-1', errors='ignore')

            # Extract printable text sequences
            import re
            # Find sequences of printable characters
            printable_parts = re.findall(r'[\x20-\x7E\n\r\t]{10,}', text)

            if printable_parts:
                return '\n'.join(printable_parts)

            raise ValueError("Could not extract readable text from DOC file")

        except Exception as e:
            logger.error(f"Failed to parse DOC file: {e}")
            raise ValueError(
                f"Could not parse DOC file. For best results, install 'antiword' "
                f"(apt-get install antiword) or convert to DOCX format. Error: {e}"
            )

    async def _parse_txt(self, file_path: Path) -> str:
        """Read plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove non-printable characters
        text = re.sub(r'[^\x20-\x7E\n]', '', text)

        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip()

    def _detect_sections(self, text: str) -> Dict[str, str]:
        """Detect and extract resume sections"""
        sections = {}
        lines = text.split('\n')

        current_section = "header"
        current_content = []

        for line in lines:
            line_stripped = line.strip()

            # Check if this line is a section header
            detected_section = None
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.match(pattern, line_stripped):
                    detected_section = section_name
                    break

            if detected_section:
                # Save previous section
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()

                current_section = detected_section
                current_content = []
            else:
                current_content.append(line)

        # Save last section
        if current_content:
            sections[current_section] = "\n".join(current_content).strip()

        return sections

    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from resume"""
        contact = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "website": None,
            "location": None
        }

        # Email pattern
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        if email_match:
            contact["email"] = email_match.group()

        # Phone pattern (various formats)
        phone_match = re.search(r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,3}[)]?[-\s\.]?[0-9]{3,4}[-\s\.]?[0-9]{3,4}', text)
        if phone_match:
            contact["phone"] = phone_match.group()

        # LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
        if linkedin_match:
            contact["linkedin"] = linkedin_match.group()

        # GitHub
        github_match = re.search(r'github\.com/[\w-]+', text, re.IGNORECASE)
        if github_match:
            contact["github"] = github_match.group()

        return contact


class ResumeSectionExtractor:
    """Helper class for more detailed section extraction"""

    @staticmethod
    def extract_experience_entries(experience_text: str) -> List[Dict]:
        """Extract individual experience entries"""
        entries = []

        # Pattern for common date formats
        date_pattern = r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}\s*[-–]\s*(?:Present|Current|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\b)'

        # Split by date patterns
        parts = re.split(date_pattern, experience_text)

        current_entry = {}
        for i, part in enumerate(parts):
            if re.match(date_pattern, part):
                if current_entry:
                    entries.append(current_entry)
                current_entry = {"date_range": part.strip()}
            elif current_entry:
                current_entry["content"] = part.strip()

        if current_entry and "content" in current_entry:
            entries.append(current_entry)

        return entries

    @staticmethod
    def extract_skills_list(skills_text: str) -> List[str]:
        """Extract individual skills from skills section"""
        # Common separators
        separators = [',', '|', '•', '·', '-', '\n']

        skills = [skills_text]
        for sep in separators:
            new_skills = []
            for skill in skills:
                new_skills.extend(skill.split(sep))
            skills = new_skills

        # Clean and filter
        skills = [s.strip() for s in skills if s.strip() and len(s.strip()) > 1]

        return skills


class ResumeChunker:
    """
    Semantic chunking for RAG (Retrieval-Augmented Generation).

    Creates meaningful chunks from resumes that can be:
    1. Stored in a vector database
    2. Retrieved during interviews for fact-checking
    3. Used for targeted question generation
    """

    # Chunk types for classification
    CHUNK_TYPES = [
        "contact_info",
        "summary",
        "experience",
        "project",
        "education",
        "skill",
        "certification",
        "achievement"
    ]

    def __init__(self, max_chunk_size: int = 500, overlap: int = 50):
        """
        Args:
            max_chunk_size: Maximum characters per chunk
            overlap: Character overlap between chunks for context
        """
        self.max_chunk_size = max_chunk_size
        self.overlap = overlap

    def chunk_resume(self, parsed_resume: Dict) -> List[Dict]:
        """
        Create semantic chunks from a parsed resume.

        Args:
            parsed_resume: Output from ResumeParser.parse()

        Returns:
            List of chunks, each with:
            - id: Unique chunk identifier
            - type: Chunk type (experience, project, skill, etc.)
            - content: Chunk text content
            - metadata: Additional context
        """
        chunks = []
        chunk_id = 0

        sections = parsed_resume.get("sections", {})
        contact_info = parsed_resume.get("contact_info", {})

        # 1. Contact info chunk
        if contact_info:
            contact_text = self._format_contact(contact_info)
            if contact_text:
                chunks.append({
                    "id": f"chunk_{chunk_id}",
                    "type": "contact_info",
                    "content": contact_text,
                    "metadata": {"section": "contact"}
                })
                chunk_id += 1

        # 2. Summary chunk
        summary = sections.get("summary", "")
        if summary:
            chunks.append({
                "id": f"chunk_{chunk_id}",
                "type": "summary",
                "content": summary[:self.max_chunk_size],
                "metadata": {"section": "summary"}
            })
            chunk_id += 1

        # 3. Experience chunks (one per job)
        experience = sections.get("experience", "")
        if experience:
            exp_chunks = self._chunk_experience(experience, chunk_id)
            chunks.extend(exp_chunks)
            chunk_id += len(exp_chunks)

        # 4. Project chunks
        projects = sections.get("projects", "")
        if projects:
            proj_chunks = self._chunk_projects(projects, chunk_id)
            chunks.extend(proj_chunks)
            chunk_id += len(proj_chunks)

        # 5. Education chunks
        education = sections.get("education", "")
        if education:
            chunks.append({
                "id": f"chunk_{chunk_id}",
                "type": "education",
                "content": education[:self.max_chunk_size],
                "metadata": {"section": "education"}
            })
            chunk_id += 1

        # 6. Skills chunk (keep together for context)
        skills = sections.get("skills", "")
        if skills:
            chunks.append({
                "id": f"chunk_{chunk_id}",
                "type": "skill",
                "content": skills[:self.max_chunk_size],
                "metadata": {
                    "section": "skills",
                    "skills_list": ResumeSectionExtractor.extract_skills_list(skills)
                }
            })
            chunk_id += 1

        # 7. Certifications chunk
        certs = sections.get("certifications", "")
        if certs:
            chunks.append({
                "id": f"chunk_{chunk_id}",
                "type": "certification",
                "content": certs[:self.max_chunk_size],
                "metadata": {"section": "certifications"}
            })
            chunk_id += 1

        # 8. Awards/achievements chunk
        awards = sections.get("awards", "")
        if awards:
            chunks.append({
                "id": f"chunk_{chunk_id}",
                "type": "achievement",
                "content": awards[:self.max_chunk_size],
                "metadata": {"section": "awards"}
            })
            chunk_id += 1

        return chunks

    def _chunk_experience(self, experience_text: str, start_id: int) -> List[Dict]:
        """Split experience section into individual job chunks."""
        chunks = []

        # Try to detect job entries by common patterns
        # Pattern: Company name followed by date range
        job_patterns = [
            # Company | Role | Date
            r'\n(?=[A-Z][A-Za-z\s&,\.]+(?:\||–|-)\s*[A-Z][a-z]+)',
            # Bullet points or new paragraphs
            r'\n\n(?=[A-Z])',
            # Date patterns as separators
            r'(?=\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\s*[-–])'
        ]

        # Try each pattern
        entries = [experience_text]
        for pattern in job_patterns:
            if len(entries) == 1:
                entries = re.split(pattern, experience_text)
                entries = [e.strip() for e in entries if e.strip()]

        # If we found distinct entries, create chunks
        if len(entries) > 1:
            for i, entry in enumerate(entries):
                if len(entry) > 50:  # Skip very short fragments
                    # Extract company/role if possible
                    first_line = entry.split('\n')[0]
                    chunks.append({
                        "id": f"chunk_{start_id + i}",
                        "type": "experience",
                        "content": entry[:self.max_chunk_size],
                        "metadata": {
                            "section": "experience",
                            "entry_index": i,
                            "title": first_line[:100]
                        }
                    })
        else:
            # Fallback: chunk by size with overlap
            chunks.extend(self._chunk_by_size(
                experience_text, "experience", start_id
            ))

        return chunks

    def _chunk_projects(self, projects_text: str, start_id: int) -> List[Dict]:
        """Split projects section into individual project chunks."""
        chunks = []

        # Try to detect project entries
        # Common patterns: numbered lists, bullet points, or headers
        project_patterns = [
            r'\n(?=\d+[\.\)]\s+)',  # Numbered list
            r'\n(?=•\s+[A-Z])',  # Bullet points
            r'\n(?=Project\s*:)',  # "Project:" prefix
            r'\n\n(?=[A-Z])',  # Double newline + capital letter
        ]

        entries = [projects_text]
        for pattern in project_patterns:
            if len(entries) == 1:
                entries = re.split(pattern, projects_text)
                entries = [e.strip() for e in entries if e.strip()]

        if len(entries) > 1:
            for i, entry in enumerate(entries):
                if len(entry) > 30:
                    # Try to extract project name
                    first_line = entry.split('\n')[0]
                    # Remove common prefixes
                    project_name = re.sub(r'^[\d\.\)\•\-\s]+', '', first_line)

                    chunks.append({
                        "id": f"chunk_{start_id + i}",
                        "type": "project",
                        "content": entry[:self.max_chunk_size],
                        "metadata": {
                            "section": "projects",
                            "project_index": i,
                            "project_name": project_name[:100],
                            "technologies": self._extract_technologies(entry)
                        }
                    })
        else:
            chunks.extend(self._chunk_by_size(
                projects_text, "project", start_id
            ))

        return chunks

    def _chunk_by_size(
        self,
        text: str,
        chunk_type: str,
        start_id: int
    ) -> List[Dict]:
        """Fallback: chunk by size with overlap."""
        chunks = []
        start = 0
        chunk_num = 0

        while start < len(text):
            end = start + self.max_chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end within last 100 chars
                last_period = text.rfind('.', end - 100, end)
                if last_period > start:
                    end = last_period + 1

            chunk_content = text[start:end].strip()

            if chunk_content:
                chunks.append({
                    "id": f"chunk_{start_id + chunk_num}",
                    "type": chunk_type,
                    "content": chunk_content,
                    "metadata": {
                        "section": chunk_type,
                        "chunk_index": chunk_num
                    }
                })
                chunk_num += 1

            start = end - self.overlap

        return chunks

    def _format_contact(self, contact_info: Dict) -> str:
        """Format contact info as text."""
        parts = []
        if contact_info.get("email"):
            parts.append(f"Email: {contact_info['email']}")
        if contact_info.get("phone"):
            parts.append(f"Phone: {contact_info['phone']}")
        if contact_info.get("linkedin"):
            parts.append(f"LinkedIn: {contact_info['linkedin']}")
        if contact_info.get("github"):
            parts.append(f"GitHub: {contact_info['github']}")
        return " | ".join(parts)

    def _extract_technologies(self, text: str) -> List[str]:
        """Extract technology mentions from project text."""
        # Common technologies to look for
        tech_patterns = [
            r'\b(Python|Java|JavaScript|TypeScript|C\+\+|C#|Go|Rust|Ruby|PHP|Swift|Kotlin)\b',
            r'\b(React|Angular|Vue|Node\.js|Django|Flask|FastAPI|Spring|Rails|Express)\b',
            r'\b(AWS|Azure|GCP|Docker|Kubernetes|Terraform|Jenkins|CI/CD)\b',
            r'\b(PostgreSQL|MySQL|MongoDB|Redis|Elasticsearch|DynamoDB|Firebase)\b',
            r'\b(TensorFlow|PyTorch|Scikit-learn|Pandas|NumPy|Keras)\b',
            r'\b(REST|GraphQL|gRPC|WebSocket|API)\b',
        ]

        technologies = set()
        for pattern in tech_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            technologies.update([m.strip() for m in matches])

        return list(technologies)

    def get_chunk_for_topic(
        self,
        chunks: List[Dict],
        topic: str,
        chunk_type: Optional[str] = None
    ) -> Optional[Dict]:
        """
        Find the most relevant chunk for a given topic/question.

        Args:
            chunks: List of resume chunks
            topic: Topic or keyword to search for
            chunk_type: Optional filter by chunk type

        Returns:
            Most relevant chunk or None
        """
        topic_lower = topic.lower()
        best_match = None
        best_score = 0

        for chunk in chunks:
            # Filter by type if specified
            if chunk_type and chunk["type"] != chunk_type:
                continue

            content_lower = chunk["content"].lower()

            # Simple relevance scoring
            score = 0

            # Exact match bonus
            if topic_lower in content_lower:
                score += 10

            # Word overlap
            topic_words = set(topic_lower.split())
            content_words = set(content_lower.split())
            overlap = len(topic_words.intersection(content_words))
            score += overlap * 2

            # Check metadata
            metadata = chunk.get("metadata", {})
            if topic_lower in str(metadata).lower():
                score += 5

            if score > best_score:
                best_score = score
                best_match = chunk

        return best_match if best_score > 0 else None
